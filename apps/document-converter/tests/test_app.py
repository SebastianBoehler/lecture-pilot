import hashlib
from io import BytesIO
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from fastapi.testclient import TestClient

from lecturepilot_converter.app import app


def test_convert_returns_revision_bound_normalized_archive(tmp_path: Path) -> None:
    source = tmp_path / "overview.docx"
    document = Document()
    document.add_heading("Course overview", level=1)
    document.add_paragraph("Evidence-based learning.")
    document.save(source)
    payload = source.read_bytes()
    sha256 = hashlib.sha256(payload).hexdigest()

    response = TestClient(app).post(
        "/convert",
        data={"source_path": "week-01/overview.docx", "source_sha256": sha256},
        files={"file": ("overview.docx", payload, "application/octet-stream")},
    )

    assert response.status_code == 200
    with ZipFile(BytesIO(response.content)) as archive:
        assert archive.namelist() == ["content.md", "manifest.json"]
        manifest = json.loads(archive.read("manifest.json"))
    assert manifest["source_sha256"] == sha256
    assert manifest["source_path"] == "week-01/overview.docx"


def test_convert_rejects_payload_that_does_not_match_revision_hash() -> None:
    response = TestClient(app).post(
        "/convert",
        data={"source_path": "week-01/overview.docx", "source_sha256": "b" * 64},
        files={
            "file": ("overview.docx", b"not the requested revision", "application/octet-stream")
        },
    )

    assert response.status_code == 400
    assert response.json() == {"detail": "Uploaded document does not match its source revision."}


def test_convert_rejects_traversal_before_processing() -> None:
    payload = b"not processed"
    response = TestClient(app).post(
        "/convert",
        data={
            "source_path": "../overview.docx",
            "source_sha256": hashlib.sha256(payload).hexdigest(),
        },
        files={"file": ("overview.docx", payload, "application/octet-stream")},
    )

    assert response.status_code == 400
    assert response.json() == {"detail": "Document path is not accepted."}


def test_convert_rejects_macro_enabled_office_documents() -> None:
    payload = b"not processed"
    response = TestClient(app).post(
        "/convert",
        data={
            "source_path": "slides/lecture.pptm",
            "source_sha256": hashlib.sha256(payload).hexdigest(),
        },
        files={"file": ("lecture.pptm", payload, "application/octet-stream")},
    )

    assert response.status_code == 400
    assert response.json() == {"detail": "Document format is not accepted."}


def test_convert_reports_malformed_office_document_without_internal_error() -> None:
    payload = b"not an OOXML document"
    response = TestClient(app).post(
        "/convert",
        data={
            "source_path": "slides/lecture.pptx",
            "source_sha256": hashlib.sha256(payload).hexdigest(),
        },
        files={"file": ("lecture.pptx", payload, "application/octet-stream")},
    )

    assert response.status_code == 422
    assert response.json() == {"detail": "Document contents could not be converted safely."}


def test_ocr_page_warns_without_blocking_when_worker_is_unavailable(monkeypatch) -> None:
    monkeypatch.delenv("LECTUREPILOT_OCR_URL", raising=False)

    response = TestClient(app).post(
        "/ocr-page",
        data={
            "native_text": "",
            "raster_ratio": "1",
            "page": "2",
            "width": "1024",
            "height": "768",
        },
        files={"file": ("page.png", b"\x89PNG\r\n\x1a\nscan", "image/png")},
    )

    assert response.status_code == 200
    assert response.json()["required"] is True
    assert response.json()["extraction"] == "native"
    assert response.json()["warning"].startswith("OCR required but unavailable for page 2")
